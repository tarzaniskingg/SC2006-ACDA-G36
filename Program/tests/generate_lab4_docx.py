"""Generate Lab4_TestCaseDesign.docx with same styling as SC2006_Lab4_v2.docx."""

from docx import Document
from docx.shared import Pt
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from lxml import etree


def set_cell_shading(cell, color):
    tc_pr = cell._element.get_or_add_tcPr()
    shading_elem = tc_pr.find(qn("w:shd"))
    if shading_elem is None:
        shading_elem = etree.SubElement(tc_pr, qn("w:shd"))
    shading_elem.set(qn("w:fill"), color)
    shading_elem.set(qn("w:val"), "clear")


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        run.font.name = "Times New Roman"
        set_cell_shading(cell, "D9E2F3")
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            run = cell.paragraphs[0].add_run(str(val))
            run.font.size = Pt(9)
            run.font.name = "Times New Roman"
    return table


def add_code(doc, code):
    for line in code.split("\n"):
        p = doc.add_paragraph()
        run = p.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(8.5)
        pf = p.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.line_spacing = Pt(12)


def body(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    run.bold = bold
    return p


def labeled(doc, label, value, mono=False):
    p = doc.add_paragraph()
    r1 = p.add_run(label)
    r1.bold = True
    r1.font.name = "Times New Roman"
    r1.font.size = Pt(12)
    r2 = p.add_run(value)
    r2.font.name = "Consolas" if mono else "Times New Roman"
    r2.font.size = Pt(10) if mono else Pt(12)


def bullet(doc, text):
    p = doc.add_paragraph(text, style="List Bullet")
    for r in p.runs:
        r.font.name = "Times New Roman"
        r.font.size = Pt(12)


def page_break(doc):
    doc.add_paragraph().add_run().add_break()


def step(doc, title, desc):
    p = doc.add_paragraph()
    r = p.add_run(title)
    r.bold = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(12)
    body(doc, desc)


# ================================================================
# Load template for styles, clear content
# ================================================================
doc = Document("SC2006_Lab4_v2.docx")
for p in list(doc.paragraphs):
    p._element.getparent().remove(p._element)
for t in list(doc.tables):
    t._element.getparent().remove(t._element)

# ================================================================
# TITLE PAGE
# ================================================================
doc.add_paragraph("Test Case Design", style="Title")
p = doc.add_paragraph("for", style="Title")
for r in p.runs:
    r.font.size = Pt(20)
doc.add_paragraph("SGTravelBud", style="Title")
doc.add_paragraph("Lab 4 Deliverable", style="ByLine")
doc.add_paragraph(
    "Prepared by Huynh Thao Tuong Van, Michelle Low Yuan-Ying, "
    "Ong Zhi Jie, Teng Wen Xuan, Terrence Ong",
    style="ByLine",
)
doc.add_paragraph("NTU", style="ByLine")
doc.add_paragraph("14 April 2026", style="ByLine")
page_break(doc)

# ================================================================
# 1. BLACK BOX TESTING
# ================================================================
doc.add_heading("1.\tBlack Box Testing: Equivalence Class + Boundary Value Testing", level=1)

doc.add_heading("1.1\tSystem Under Test", level=2)
labeled(doc, "Control class: ", "CostEstimationController (backend/services/routing.py)")
labeled(
    doc,
    "Function: ",
    "estimate_cost(distance_m, duration_s, mode, departure_time, origin_lat, origin_lng)",
    mono=True,
)
labeled(
    doc,
    "Description: ",
    "Calculates trip cost based on transport mode. For taxi/driving: metered fare "
    "with distance tiers, time-based surcharges, and airport surcharge. For transit: "
    "distance-based fare with time-of-day discounts. For own car: fuel cost only.",
)
body(
    doc,
    "For black box testing, we focus on taxi mode with two key range parameters: "
    "distance_m and departure_time.",
)

doc.add_heading("1.2\tRequirements Specification", level=2)
for s in [
    "The function takes distance_m, a float >= 0 (distance in meters)",
    "The function takes departure_time, a datetime (Singapore time)",
    "If distance <= 1km: flag-down fare only ($4.60)",
    "If 1km < distance <= 10km: $0.27 per 400m after first km (Tier 1)",
    "If distance > 10km: $0.27 per 350m after 10km (Tier 2)",
    "If departure is Mon-Fri 6:00am-9:29am, OR any day 5:00pm-11:59pm, "
    "OR Sat-Sun 10:00am-1:59pm: +25% peak surcharge",
    "If departure is midnight-5:59am: +50% late-night surcharge (replaces peak)",
    "Otherwise: no surcharge",
]:
    bullet(doc, s)

# 1.3 ECs
doc.add_heading("1.3\tEquivalence Classes", level=2)
body(doc, "Parameter 1: distance_m (float) \u2014 Range of values", bold=True)
add_table(
    doc,
    ["EC ID", "Range (km)", "Type", "Expected behaviour"],
    [
        ["D1", "d <= 0", "Invalid", "Clamped to 0, flag-down only"],
        ["D2", "0 < d <= 1", "Valid", "Flag-down only ($4.60 base)"],
        ["D3", "1 < d <= 10", "Valid", "Flag-down + Tier 1 rate"],
        ["D4", "d > 10", "Valid", "Flag-down + Tier 1 + Tier 2 rate"],
    ],
)

doc.add_paragraph()
body(doc, "Parameter 2: departure_time (datetime) \u2014 Range of values", bold=True)
add_table(
    doc,
    ["EC ID", "Condition", "Type", "Expected behaviour"],
    [
        ["T1", "Weekday 6:00am-9:29am", "Valid", "+25% peak surcharge, booking $3.30"],
        ["T2", "Any day 5:00pm-11:59pm", "Valid", "+25% peak surcharge, booking $3.30"],
        ["T3", "Sat-Sun 10:00am-1:59pm", "Valid", "+25% peak surcharge, booking $3.30"],
        ["T4", "Midnight-5:59am", "Valid", "+50% late-night, booking $3.30"],
        ["T5", "All other times", "Valid", "No surcharge, booking $2.30"],
    ],
)

# 1.4 BVs
doc.add_heading("1.4\tBoundary Values", level=2)
body(doc, "Parameter 1: distance_m boundary values", bold=True)
add_table(
    doc,
    ["EC boundary", "just-below", "on-boundary", "just-above"],
    [
        ["D1/D2: 0m", "-1 (D1)", "0 (D1)", "1 (D2)"],
        ["D2/D3: 1000m", "999 (D2)", "1000 (D2)", "1001 (D3)"],
        ["D3/D4: 10000m", "9999 (D3)", "10000 (D3)", "10001 (D4)"],
    ],
)

doc.add_paragraph()
body(doc, "Parameter 2: departure_time boundary values (minutes from midnight, weekday)", bold=True)
add_table(
    doc,
    ["EC boundary", "just-below", "on-boundary", "just-above"],
    [
        ["T4/T1: 360 (6:00am)", "359 (T4)", "360 (T1)", "361 (T1)"],
        ["T1/T5: 570 (9:30am)", "569 (T1)", "570 (T5)", "571 (T5)"],
        ["T5/T2: 1020 (5:00pm)", "1019 (T5)", "1020 (T2)", "1021 (T2)"],
    ],
)

# 1.5 Valid combos
doc.add_heading("1.5\tTest Cases \u2014 Valid Input Combinations", level=2)
add_table(
    doc,
    ["TC", "distance_m", "departure_time", "Expected surcharge", "Expected tier"],
    [
        ["V1", "1000", "Tue 14:00 (T5)", "No surcharge", "Flag-down only"],
        ["V2", "5000", "Tue 14:00 (T5)", "No surcharge", "Tier 1"],
        ["V3", "15000", "Tue 14:00 (T5)", "No surcharge", "Tier 1 + Tier 2"],
        ["V4", "5000", "Tue 08:00 (T1)", "+25% peak", "Tier 1"],
        ["V5", "5000", "Tue 18:00 (T2)", "+25% peak", "Tier 1"],
        ["V6", "5000", "Tue 03:00 (T4)", "+50% late-night", "Tier 1"],
    ],
)

# 1.6 Invalid
doc.add_heading("1.6\tTest Cases \u2014 Invalid Input (one invalid at a time)", level=2)
add_table(
    doc,
    ["TC", "distance_m", "departure_time", "Expected result"],
    [
        ["I1", "-1 (invalid D1)", "Tue 14:00 (valid T5)", "distance_charge=0, total>0"],
        ["I2", "0 (invalid D1)", "Tue 08:00 (valid T1)", "distance_charge=0, peak applies"],
    ],
)

# 1.7 BV tests
doc.add_heading("1.7\tTest Cases \u2014 Boundary Value Tests", level=2)
body(doc, "Distance tier boundaries:", bold=True)
add_table(
    doc,
    ["TC", "distance_m", "departure_time", "Expected"],
    [
        ["B1", "999", "Tue 14:00", "distance_charge = 0 (within flag-down)"],
        ["B2", "1000", "Tue 14:00", "distance_charge = 0 (on boundary)"],
        ["B3", "1100", "Tue 14:00", "distance_charge > 0 (Tier 1 active)"],
        ["B4", "9999", "Tue 14:00", "Tier 1 rate only"],
        ["B5", "10000", "Tue 14:00", "Tier 1 rate only (on boundary)"],
        ["B6", "10500", "Tue 14:00", "Tier 1 + Tier 2 (above boundary)"],
    ],
)

doc.add_paragraph()
body(doc, "Time surcharge boundaries:", bold=True)
add_table(
    doc,
    ["TC", "distance_m", "departure_time", "Expected"],
    [
        ["B7", "5000", "Tue 05:59 (t=359)", "Late-night surcharge (hour < 6)"],
        ["B8", "5000", "Tue 06:00 (t=360)", "Peak surcharge (morning start)"],
        ["B9", "5000", "Tue 09:29 (t=569)", "Peak surcharge (morning end)"],
        ["B10", "5000", "Tue 09:30 (t=570)", "No surcharge (just after peak)"],
        ["B11", "5000", "Tue 16:59 (t=1019)", "No surcharge (before evening)"],
        ["B12", "5000", "Tue 17:00 (t=1020)", "Peak surcharge (evening start)"],
    ],
)

# 1.8 Results
doc.add_heading("1.8\tTest Results", level=2)
add_table(
    doc,
    ["TC", "Input", "Oracle (Expected)", "Log (Actual)", "Pass?"],
    [
        ["V1", "(1000, 600, taxi, Tue 14:00)", "dist=0, peak=0, late=0, fee=2.30", "dist=0.0, peak=0.0, late=0.0, fee=2.30, total=7.62", "PASS"],
        ["V2", "(5000, 600, taxi, Tue 14:00)", "dist=2.70, peak=0, fee=2.30", "dist=2.70, peak=0.0, fee=2.30, total=10.32", "PASS"],
        ["V3", "(15000, 1200, taxi, Tue 14:00)", "dist>6.08, peak=0, fee=2.30", "dist=9.93, peak=0.0, fee=2.30, total=18.27", "PASS"],
        ["V4", "(5000, 600, taxi, Tue 08:00)", "peak>0, late=0, fee=3.30", "peak=2.0, late=0.0, fee=3.30, total=13.32", "PASS"],
        ["V5", "(5000, 600, taxi, Tue 18:00)", "peak>0, late=0, fee=3.30", "peak=2.0, late=0.0, fee=3.30, total=13.32", "PASS"],
        ["V6", "(5000, 600, taxi, Tue 03:00)", "late>0, peak=0, fee=3.30", "late=4.01, peak=0.0, fee=3.30, total=15.33", "PASS"],
        ["I1", "(-1, 600, taxi, Tue 14:00)", "dist_charge=0", "dist=0.0, total=7.62", "PASS"],
        ["I2", "(0, 600, taxi, Tue 08:00)", "dist=0, peak>0", "dist=0.0, peak=1.38", "PASS"],
        ["B1", "(999, 600, taxi, Tue 14:00)", "dist_charge=0", "dist=0.0", "PASS"],
        ["B2", "(1000, 600, taxi, Tue 14:00)", "dist_charge=0", "dist=0.0", "PASS"],
        ["B3", "(1100, 600, taxi, Tue 14:00)", "dist_charge>0", "dist=0.07", "PASS"],
        ["B4", "(9999, 600, taxi, Tue 14:00)", "Tier 1 only", "dist=6.07", "PASS"],
        ["B5", "(10000, 600, taxi, Tue 14:00)", "dist=6.08", "dist=6.08, total=13.70", "PASS"],
        ["B6", "(10500, 600, taxi, Tue 14:00)", "dist>6.08", "dist=6.46, total=14.08", "PASS"],
        ["B7", "(5000, 600, taxi, Tue 05:59)", "late_night>0", "late=4.01, total=15.33", "PASS"],
        ["B8", "(5000, 600, taxi, Tue 06:00)", "peak>0, late=0", "peak=2.0, late=0.0", "PASS"],
        ["B9", "(5000, 600, taxi, Tue 09:29)", "peak>0", "peak=2.0", "PASS"],
        ["B10", "(5000, 600, taxi, Tue 09:30)", "peak=0, late=0", "peak=0.0, late=0.0", "PASS"],
        ["B11", "(5000, 600, taxi, Tue 16:59)", "peak=0", "peak=0.0", "PASS"],
        ["B12", "(5000, 600, taxi, Tue 17:00)", "peak>0", "peak=2.0", "PASS"],
    ],
)

body(doc, "Total black box test cases: 20 (6 valid + 2 invalid + 12 boundary). All tests: PASSED.", bold=True)
page_break(doc)

# ================================================================
# 2. WHITE BOX TESTING
# ================================================================
doc.add_heading("2.\tWhite Box Testing: Basis Path Testing", level=1)

# --- 2.1 _is_peak ---
doc.add_heading("2.1\tMethod 1: _is_peak(dt)", level=2)

doc.add_heading("2.1.1\tSource Code", level=3)
labeled(doc, "Location: ", "backend/services/routing.py, line 52")
add_code(doc, """def _is_peak(dt: datetime) -> bool:
1    wd = dt.weekday()
2    h, m = dt.hour, dt.minute
3    t = h * 60 + m
4    if wd < 5 and 360 <= t <= 569:
5        return True
6    if t >= 1020:
7        return True
8    if wd >= 5 and 600 <= t <= 839:
9        return True
10   return False""")

doc.add_heading("2.1.2\tControl Flow Graph", level=3)
body(doc, "The CFG contains 7 nodes and 7 edges:")
add_code(doc, """        [1-3: compute wd, t]
               |
               v
        <D1: wd<5 AND 360<=t<=569?>
          /              \\
        True            False
         |                |
         v                v
    [5: return True]  <D2: t>=1020?>
                       /          \\
                     True        False
                      |            |
                      v            v
                [7: return    <D3: wd>=5 AND
                   True]       600<=t<=839?>
                               /          \\
                             True        False
                              |            |
                              v            v
                        [9: return    [10: return
                           True]        False]""")

doc.add_heading("2.1.3\tCyclomatic Complexity", level=3)
body(doc, "Decision points: D1, D2, D3 = 3")
body(doc, "V(G) = decision points + 1 = 3 + 1 = 4", bold=True)
body(doc, "This means we need 4 linearly independent basis paths.")

doc.add_heading("2.1.4\tBasis Path Derivation", level=3)
step(doc, "Step 1: Choose baseline path",
     "Baseline: 1-3 \u2192 D1(True) \u2192 5 (return True). Represents weekday morning peak.")
step(doc, "Step 2: Mutate D1 (flip to False)",
     "Path 2: 1-3 \u2192 D1(False) \u2192 D2(True) \u2192 7 (return True). Now takes evening peak branch.")
step(doc, "Step 3: Mutate D2 (flip to False)",
     "Path 3: 1-3 \u2192 D1(False) \u2192 D2(False) \u2192 D3(True) \u2192 9 (return True). Now takes weekend midday branch.")
step(doc, "Step 4: Mutate D3 (flip to False)",
     "Path 4: 1-3 \u2192 D1(False) \u2192 D2(False) \u2192 D3(False) \u2192 10 (return False). All decisions false = off-peak.")

doc.add_heading("2.1.5\tTest Cases and Results", level=3)
add_table(
    doc,
    ["Path", "Node sequence", "Input", "Oracle", "Log", "Pass?"],
    [
        ["1", "1-3, D1(T), 5", "Tue 08:00 SGT (wd=1, t=480)", "True", "True", "PASS"],
        ["2", "1-3, D1(F), D2(T), 7", "Tue 18:00 SGT (wd=1, t=1080)", "True", "True", "PASS"],
        ["3", "1-3, D1(F), D2(F), D3(T), 9", "Sat 12:00 SGT (wd=5, t=720)", "True", "True", "PASS"],
        ["4", "1-3, D1(F), D2(F), D3(F), 10", "Tue 14:00 SGT (wd=1, t=840)", "False", "False", "PASS"],
    ],
)
body(doc, "Coverage achieved: 100% statement, 100% branch, 100% basis path.", bold=True)
page_break(doc)

# --- 2.2 _time_based_crowding ---
doc.add_heading("2.2\tMethod 2: _time_based_crowding(query_time)", level=2)

doc.add_heading("2.2.1\tSource Code", level=3)
labeled(doc, "Location: ", "backend/services/assessment.py, line 227")
add_code(doc, """def _time_based_crowding(query_time=None) -> str:
1    SGT = timezone(timedelta(hours=8))
2    dt = query_time or datetime.now(SGT)
3    if dt.tzinfo is None:
4        dt = dt.replace(tzinfo=SGT)
     else:
5        dt = dt.astimezone(SGT)
6    h, m = dt.hour, dt.minute
7    t = h * 60 + m
8    wd = dt.weekday()
9    if wd < 5:                         # Weekday
10       if 420 <= t <= 570:            #   Morning peak
11           return "High"
12       if 1050 <= t <= 1200:          #   Evening peak
13           return "High"
14       if (360<=t<420) or (570<t<=630)
              or (990<=t<1050) or (1200<t<=1260):
15           return "Medium"            #   Shoulder
16       return "Low"                   #   Off-peak
     else:                              # Weekend
17       if 660 <= t <= 1080:           #   Midday busy
18           return "Medium"
19       return "Low" """)

doc.add_heading("2.2.2\tControl Flow Graph", level=3)
add_code(doc, """        [1-8: compute dt, t, wd]
               |
               v
        <D1: wd < 5?>
          /            \\
        True          False
         |               |
         v               v
    <D2: 420<=t<=570?>  <D5: 660<=t<=1080?>
      /        \\          /          \\
    True      False     True        False
     |          |        |            |
     v          v        v            v
 [11: ret   <D3: 1050  [18: ret   [19: ret
  "High"]   <=t<=1200?>  "Med"]     "Low"]
              /        \\
            True      False
             |          |
             v          v
         [13: ret   <D4: shoulder?>
          "High"]     /        \\
                    True      False
                     |          |
                     v          v
                 [15: ret   [16: ret
                  "Med"]     "Low"]""")

doc.add_heading("2.2.3\tCyclomatic Complexity", level=3)
body(doc, "Decision points: D1, D2, D3, D4, D5 = 5")
body(doc, "V(G) = decision points + 1 = 5 + 1 = 6", bold=True)
body(doc, "This means we need 6 linearly independent basis paths.")

doc.add_heading("2.2.4\tBasis Path Derivation", level=3)
step(doc, "Step 1: Choose baseline path",
     'Baseline: 1-8 \u2192 D1(T) \u2192 D2(T) \u2192 11 (return "High"). Represents weekday morning peak.')
step(doc, "Step 2: Mutate D2 (flip to False)",
     'Path 2: 1-8 \u2192 D1(T) \u2192 D2(F) \u2192 D3(T) \u2192 13 (return "High"). Now takes evening peak.')
step(doc, "Step 3: Mutate D3 (flip to False)",
     'Path 3: 1-8 \u2192 D1(T) \u2192 D2(F) \u2192 D3(F) \u2192 D4(T) \u2192 15 (return "Medium"). Shoulder period.')
step(doc, "Step 4: Mutate D4 (flip to False)",
     'Path 4: 1-8 \u2192 D1(T) \u2192 D2(F) \u2192 D3(F) \u2192 D4(F) \u2192 16 (return "Low"). Weekday off-peak.')
step(doc, "Step 5: Mutate D1 (flip to False)",
     'Path 5: 1-8 \u2192 D1(F) \u2192 D5(T) \u2192 18 (return "Medium"). Weekend midday.')
step(doc, "Step 6: Mutate D5 (flip to False)",
     'Path 6: 1-8 \u2192 D1(F) \u2192 D5(F) \u2192 19 (return "Low"). Weekend off-peak.')

doc.add_heading("2.2.5\tTest Cases and Results", level=3)
add_table(
    doc,
    ["Path", "Node sequence", "Input", "Oracle", "Log", "Pass?"],
    [
        ["1", "1-8, D1(T), D2(T), 11", "Tue 08:00 (wd=1, t=480)", "High", "High", "PASS"],
        ["2", "1-8, D1(T), D2(F), D3(T), 13", "Tue 18:20 (wd=1, t=1100)", "High", "High", "PASS"],
        ["3", "1-8, D1(T), D2(F), D3(F), D4(T), 15", "Tue 10:00 (wd=1, t=600)", "Medium", "Medium", "PASS"],
        ["4", "1-8, D1(T), D2(F), D3(F), D4(F), 16", "Tue 14:00 (wd=1, t=840)", "Low", "Low", "PASS"],
        ["5", "1-8, D1(F), D5(T), 18", "Sat 12:00 (wd=5, t=720)", "Medium", "Medium", "PASS"],
        ["6", "1-8, D1(F), D5(F), 19", "Sat 09:00 (wd=5, t=540)", "Low", "Low", "PASS"],
    ],
)
body(doc, "Coverage achieved: 100% statement, 100% branch, 100% basis path.", bold=True)
page_break(doc)

# ================================================================
# 3. SUMMARY
# ================================================================
doc.add_heading("3.\tSummary", level=1)
add_table(
    doc,
    ["Technique", "System Under Test", "Test cases", "Result"],
    [
        ["Black box (EC + BVT)", "estimate_cost() \u2014 taxi mode", "20", "All PASS"],
        ["White box (basis path)", "_is_peak() \u2014 CC=4", "4", "All PASS"],
        ["White box (basis path)", "_time_based_crowding() \u2014 CC=6", "6", "All PASS"],
        ["Total", "", "30", "All PASS"],
    ],
)
doc.add_paragraph()
body(doc, "Automated test files (96 tests including additional coverage):")
bullet(doc, "tests/test_blackbox_estimate_cost.py (60 automated tests)")
bullet(doc, "tests/test_whitebox_is_peak.py (11 automated tests)")
bullet(doc, "tests/test_whitebox_time_based_crowding.py (25 automated tests)")
doc.add_paragraph()
labeled(doc, "Run command: ", 'python -m unittest discover -s tests -p "test_*.py" -v', mono=True)

# ================================================================
doc.save("tests/Lab4_TestCaseDesign.docx")
print("DONE: tests/Lab4_TestCaseDesign.docx")
