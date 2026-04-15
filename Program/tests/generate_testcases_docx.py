"""Generate TestCases.docx from structured test case data."""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = Document()

# -- Styles --
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(10)
style.paragraph_format.space_after = Pt(4)

# Narrow margins
for section in doc.sections:
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

# -- Title --
title = doc.add_heading("SGTravelBud — Test Cases", level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph(
    "Comprehensive test cases covering all functionality of the SGTravelBud application."
)

# -- Test case data --
test_cases = [
    {
        "id": "TC1",
        "title": "Route Search — Basic Transit Search",
        "objective": "Verify that a basic public transit route search returns valid results.",
        "preconditions": "Backend server running; Google Maps API key configured.",
        "steps": (
            "1. Open the app (MainView).\n"
            "2. Enter \"Clementi MRT\" as origin.\n"
            "3. Enter \"Raffles Place MRT\" as destination.\n"
            "4. Ensure \"Public Transit\" mode is selected.\n"
            "5. Tap Search."
        ),
        "expected": (
            "Bottom sheet expands showing 1\u20135 route cards. Each card displays duration, "
            "cost, transfer count, and a crowding risk badge (Low/Medium/High). "
            "Map updates with route polyline."
        ),
    },
    {
        "id": "TC2",
        "title": "Route Search — Driving / Taxi Mode",
        "objective": "Verify driving and taxi routes are returned when the Taxi/Drive toggle is active.",
        "preconditions": "Backend running.",
        "steps": (
            "1. Enter origin \"NUS\" and destination \"Changi Airport\".\n"
            "2. Toggle mode to Taxi/Drive (deselect Public Transit).\n"
            "3. Tap Search."
        ),
        "expected": (
            "Results contain route cards categorised as \"Taxi\" or \"Drive\". "
            "Cost reflects estimated taxi fare or fuel cost. "
            "No transit-specific fields (transfers) shown."
        ),
    },
    {
        "id": "TC3",
        "title": "Route Search — Both Modes Selected",
        "objective": "Verify mixed-mode results when both Public Transit and Taxi/Drive are enabled.",
        "preconditions": "Backend running.",
        "steps": (
            "1. Enter origin and destination.\n"
            "2. Enable both mode toggles.\n"
            "3. Tap Search."
        ),
        "expected": (
            "Results include a mix of transit and driving/taxi routes, deduplicated, "
            "up to 5 total. Categories are labelled correctly."
        ),
    },
    {
        "id": "TC4",
        "title": "Route Search — Swap Origin and Destination",
        "objective": "Verify the swap button correctly reverses origin and destination.",
        "preconditions": "Backend running.",
        "steps": (
            "1. Enter \"Jurong East\" as origin and \"Marina Bay\" as destination.\n"
            "2. Tap the Swap button."
        ),
        "expected": "Origin becomes \"Marina Bay\" and destination becomes \"Jurong East\".",
    },
    {
        "id": "TC5",
        "title": "Route Search — Validation (Empty Fields)",
        "objective": "Verify search is blocked when required fields are incomplete.",
        "preconditions": "App open on MainView.",
        "steps": (
            "1. Leave origin empty, enter destination.\n"
            "2. Attempt to tap Search."
        ),
        "expected": "Search button is disabled. No API call is made.",
    },
    {
        "id": "TC6",
        "title": "Route Search — Validation (Same Origin and Destination)",
        "objective": "Verify search rejects identical origin and destination.",
        "preconditions": "App open on MainView.",
        "steps": (
            "1. Enter \"Orchard MRT\" as both origin and destination.\n"
            "2. Tap Search."
        ),
        "expected": "Search button is disabled or an error message is shown. No results returned.",
    },
    {
        "id": "TC7",
        "title": "Route Search — No Routes Found",
        "objective": "Verify graceful handling when no routes exist for the query.",
        "preconditions": "Backend running.",
        "steps": (
            "1. Enter a very remote or invalid location pair.\n"
            "2. Tap Search."
        ),
        "expected": "\"No routes found\" message displayed. App does not crash.",
    },
    {
        "id": "TC8",
        "title": "Advanced Priorities — Weight Sliders",
        "objective": "Verify that adjusting weight sliders changes route ranking.",
        "preconditions": "Backend running.",
        "steps": (
            "1. Enter origin/destination and search with default weights (all 0.25).\n"
            "2. Note the top-ranked route.\n"
            "3. Expand Advanced Priorities.\n"
            "4. Set Time weight to 1.0 and all others to 0.0.\n"
            "5. Search again."
        ),
        "expected": (
            "The fastest route is now ranked first. Route ordering changes compared to step 2."
        ),
    },
    {
        "id": "TC9",
        "title": "Constraints — Max Walk Time",
        "objective": "Verify max walk constraint filters routes appropriately.",
        "preconditions": "Backend running.",
        "steps": (
            "1. Expand Advanced Priorities.\n"
            "2. Set Max Walk to 5 minutes.\n"
            "3. Search for a route that normally involves long walks."
        ),
        "expected": "Routes with walking segments exceeding 5 minutes are excluded from results.",
    },
    {
        "id": "TC10",
        "title": "Constraints — Max Transfers",
        "objective": "Verify max transfers constraint filters routes.",
        "preconditions": "Backend running.",
        "steps": (
            "1. Set Max Transfers to 0.\n"
            "2. Search for a multi-transfer route."
        ),
        "expected": (
            "Only direct (no-transfer) routes are returned. "
            "If none exist, \"No routes found\" is shown."
        ),
    },
    {
        "id": "TC11",
        "title": "Constraints — Max Budget",
        "objective": "Verify max budget constraint filters routes by cost.",
        "preconditions": "Backend running.",
        "steps": (
            "1. Set Max Budget to $1.00.\n"
            "2. Search for a route."
        ),
        "expected": "Only routes costing $1.00 or less are returned. Expensive routes are excluded.",
    },
    {
        "id": "TC12",
        "title": "Results Page — Route Selection and Map Navigation",
        "objective": "Verify selecting a route and navigating to map view.",
        "preconditions": "A successful search has been performed.",
        "steps": (
            "1. On ResultsPage, tap a RouteCard.\n"
            "2. Tap \"View on Map\"."
        ),
        "expected": (
            "MapPage opens showing the selected route on an interactive map with "
            "start/end markers and a polyline. Bottom panel shows time, distance, cost, risk badges."
        ),
    },
    {
        "id": "TC13",
        "title": "Map Page — Navigate Between Routes",
        "objective": "Verify Prev/Next buttons cycle through routes on the map.",
        "preconditions": "MapPage open with multiple routes.",
        "steps": (
            "1. Tap Next to go to route 2.\n"
            "2. Tap Prev to return to route 1."
        ),
        "expected": (
            "Map updates to show the corresponding route. \"Route X of Y\" label updates. "
            "Prev is disabled on route 1; Next is disabled on the last route."
        ),
    },
    {
        "id": "TC14",
        "title": "Scoring Page — Breakdown Display",
        "objective": "Verify scoring breakdown is displayed correctly for multiple routes.",
        "preconditions": "A search returning 2+ routes has been performed.",
        "steps": "1. Navigate to the Scoring tab.",
        "expected": (
            "Page shows: active weights, comparison table (best values green, worst red), "
            "normalised bar charts per route, weighted contribution grid, and formula breakdown."
        ),
    },
    {
        "id": "TC15",
        "title": "Scoring Page — Single Route",
        "objective": "Verify scoring page handles a single route gracefully.",
        "preconditions": "A search returning exactly 1 route.",
        "steps": "1. Navigate to the Scoring tab.",
        "expected": (
            "Comparison table is hidden (not applicable). "
            "Single route breakdown card is shown with rank \"Best\"."
        ),
    },
    {
        "id": "TC16",
        "title": "Departure Time Comparison — Display",
        "objective": "Verify the time comparison modal shows routes across 9 time slots.",
        "preconditions": "A successful search has been performed.",
        "steps": (
            "1. On ResultsPage, tap Compare Departure Times."
        ),
        "expected": (
            "Modal opens with 3 groups (Morning Rush, Around Now, Evening Rush). "
            "Each group has 3 time slots. Cells show score, time, realistic time, cost, and crowding. "
            "Best values highlighted green, worst red."
        ),
    },
    {
        "id": "TC17",
        "title": "Departure Time Comparison — Close Modal",
        "objective": "Verify the modal can be dismissed.",
        "preconditions": "Time comparison modal is open.",
        "steps": (
            "1. Tap the close button (or click backdrop)."
        ),
        "expected": "Modal closes. Results page is visible again.",
    },
    {
        "id": "TC18",
        "title": "Crowding Heatmap — Display",
        "objective": "Verify crowding heatmap renders for train stations.",
        "preconditions": "A transit search involving MRT stations has been performed.",
        "steps": "1. Navigate to Scoring page.",
        "expected": (
            "Heatmap bars appear for each train station on the route. "
            "Color blocks show Low (green), Medium (amber), High (red) for each hour from 6 AM to 11 PM. "
            "A \"Now\" marker indicates the current time."
        ),
    },
    {
        "id": "TC19",
        "title": "Settings Page — Load and Display",
        "objective": "Verify settings page loads saved preferences.",
        "preconditions": "Backend running.",
        "steps": "1. Navigate to the Settings tab.",
        "expected": (
            "Page shows current language, units, and default weight sliders loaded from backend. "
            "Data sources table shows status (Live/Fallback) for each dataset."
        ),
    },
    {
        "id": "TC20",
        "title": "Settings Page — Save Preferences",
        "objective": "Verify settings can be saved and persist.",
        "preconditions": "Backend running.",
        "steps": (
            "1. Navigate to Settings.\n"
            "2. Change language to \"zh\" and units to \"imperial\".\n"
            "3. Adjust default Time weight to 0.8.\n"
            "4. Tap Save.\n"
            "5. Navigate away, then return to Settings."
        ),
        "expected": (
            "Green success toast appears on save. On return, settings reflect the saved values "
            "(zh, imperial, Time=0.8)."
        ),
    },
    {
        "id": "TC21",
        "title": "Settings Page — Refresh Data Sources",
        "objective": "Verify cache refresh invalidates and re-fetches data.",
        "preconditions": "Backend running with cached data.",
        "steps": (
            "1. Navigate to Settings.\n"
            "2. Note the current data source statuses.\n"
            "3. Tap the Refresh button."
        ),
        "expected": (
            "All caches are invalidated. Data sources table reloads. "
            "Sources that successfully re-fetch show \"Live\"; those that fail show \"Fallback\"."
        ),
    },
    {
        "id": "TC22",
        "title": "Risk Badges — Correct Categorisation",
        "objective": "Verify risk badges display the correct category and colour.",
        "preconditions": "A transit search has been performed.",
        "steps": "1. Observe crowding and delay badges on route cards and MapPage.",
        "expected": (
            "Badges show \"Low\" (green), \"Medium\" (amber), or \"High\" (red) "
            "matching the backend assessment data."
        ),
    },
    {
        "id": "TC23",
        "title": "API — Route Search Endpoint",
        "objective": "Verify GET /routes returns correct response structure.",
        "preconditions": "Backend running.",
        "steps": (
            "Send: GET /routes?origin=Clementi&destination=Raffles+Place"
            "&include_transit=true&wt_time=0.5&wt_cost=0.2&wt_risk=0.2&wt_comfort=0.1"
        ),
        "expected": (
            "200 OK. Response contains \"trip\" (origin, destination, weights), "
            "\"routes\" (list of up to 5 RouteOption objects with time, cost, score, "
            "steps, crowding, delay fields)."
        ),
    },
    {
        "id": "TC24",
        "title": "API — Route Search with Constraints",
        "objective": "Verify constraint parameters filter results server-side.",
        "preconditions": "Backend running.",
        "steps": "Send: GET /routes?origin=NUS&destination=Changi+Airport&max_transfers=1&max_budget=3.0",
        "expected": "200 OK. All returned routes have transfers <= 1 and cost <= $3.00.",
    },
    {
        "id": "TC25",
        "title": "API — Assessment Endpoint",
        "objective": "Verify GET /assessment returns segment-level risk data.",
        "preconditions": "Backend running.",
        "steps": "Send: GET /assessment?origin=Jurong+East&destination=Dhoby+Ghaut",
        "expected": (
            "200 OK. Response contains segment assessments with crowding levels, "
            "delay estimates, and bus frequency data per segment."
        ),
    },
    {
        "id": "TC26",
        "title": "API — Crowding Heatmap Endpoint",
        "objective": "Verify GET /crowding/heatmap returns hourly crowding data.",
        "preconditions": "Backend running.",
        "steps": "Send: GET /crowding/heatmap?station_name=Orchard",
        "expected": (
            "200 OK. Response contains station name, line code, and intervals array "
            "with entries from 6 AM to 11 PM, each having a crowding level (l/m/h)."
        ),
    },
    {
        "id": "TC27",
        "title": "API — Settings Round-Trip",
        "objective": "Verify settings can be saved and loaded via API.",
        "preconditions": "Backend running.",
        "steps": (
            "1. Send PUT /settings with body: {\"language\":\"ta\", \"units\":\"imperial\", "
            "\"default_weights\":{\"time\":0.4, \"cost\":0.3, \"risk\":0.2, \"comfort\":0.1}}.\n"
            "2. Send GET /settings."
        ),
        "expected": "PUT returns 200 with updated settings. GET returns the same values that were saved.",
    },
    {
        "id": "TC28",
        "title": "API — Cache Status and Refresh",
        "objective": "Verify dataset status reporting and cache invalidation.",
        "preconditions": "Backend running with cached data.",
        "steps": (
            "1. Send GET /datasets \u2014 note statuses.\n"
            "2. Send POST /refresh.\n"
            "3. Send GET /datasets again."
        ),
        "expected": (
            "Initial GET shows cached data with timestamps. POST invalidates all caches. "
            "Second GET shows refreshed timestamps. Fallback flags update based on API availability."
        ),
    },
    {
        "id": "TC29",
        "title": "API — Health Check",
        "objective": "Verify the health endpoint responds.",
        "preconditions": "Backend running.",
        "steps": "Send: GET /health",
        "expected": "200 OK with a health status response.",
    },
    {
        "id": "TC30",
        "title": "Scoring Logic — Weighted Composite Score",
        "objective": "Verify scoring formula produces correct ranking.",
        "preconditions": "Backend running.",
        "steps": (
            "1. Search with weights: Time=1.0, Cost=0.0, Risk=0.0, Comfort=0.0.\n"
            "2. Verify the fastest route has the lowest score.\n"
            "3. Repeat with Cost=1.0 (others 0.0).\n"
            "4. Verify the cheapest route has the lowest score."
        ),
        "expected": (
            "Route with the best value in the weighted dimension always ranks first. "
            "Score = sum of (normalised_metric \u00d7 weight)."
        ),
    },
    {
        "id": "TC31",
        "title": "Fallback Behaviour — API Failure",
        "objective": "Verify app uses cached fallback data when an external API is unavailable.",
        "preconditions": "Perform a search once (to populate cache), then simulate LTA API unavailability.",
        "steps": (
            "1. Perform a route search (cache populated).\n"
            "2. Block LTA API access.\n"
            "3. Perform the same search again."
        ),
        "expected": (
            "Results still return using cached data. Backend message may include "
            "\"Using fallback data\". Settings page shows \"Fallback\" badge for affected data source."
        ),
    },
    {
        "id": "TC32",
        "title": "ERP Cost Calculation (Driving Routes)",
        "objective": "Verify ERP gantry charges are included in driving route costs.",
        "preconditions": "Backend running.",
        "steps": (
            "1. Search a driving route that passes through known ERP gantries "
            "(e.g., CTE during peak hours).\n"
            "2. Check the cost breakdown."
        ),
        "expected": "Route cost includes ERP toll charges. Explanation text mentions ERP if applicable.",
    },
    {
        "id": "TC33",
        "title": "Weather Integration",
        "objective": "Verify weather data is factored into route assessment.",
        "preconditions": "Backend running; NEA weather API accessible.",
        "steps": (
            "1. Perform a route search.\n"
            "2. Check if weather information appears in route explanation or risk assessment."
        ),
        "expected": (
            "If adverse weather is forecast for the route area, it is reflected in the "
            "risk assessment or route explanation."
        ),
    },
    {
        "id": "TC34",
        "title": "Parking Availability (Drive Mode)",
        "objective": "Verify nearby carpark availability is shown for driving routes.",
        "preconditions": "Backend running.",
        "steps": (
            "1. Search a driving route to a destination with nearby HDB/URA carparks.\n"
            "2. Check route details."
        ),
        "expected": "Carpark availability data is included in the route response. Available lots count shown.",
    },
    {
        "id": "TC35",
        "title": "Realistic Time Calculation",
        "objective": "Verify realistic time accounts for bus frequency and delays.",
        "preconditions": "Backend running.",
        "steps": (
            "1. Search a transit route involving buses.\n"
            "2. Compare Google time vs Realistic time on the Scoring page."
        ),
        "expected": (
            "Realistic time >= Google time. The difference reflects added delay from bus frequency risk. "
            "Realistic time is highlighted amber on the Scoring page when it exceeds Google time."
        ),
    },
    {
        "id": "TC36",
        "title": "Place Autocomplete (Geocoding)",
        "objective": "Verify location autocomplete suggestions appear while typing.",
        "preconditions": "App open; internet available.",
        "steps": (
            "1. Click the origin field.\n"
            "2. Type \"Clem\" (at least 2 characters)."
        ),
        "expected": (
            "Dropdown suggestions appear from Nominatim API "
            "(e.g., \"Clementi MRT Station\", \"Clementi Avenue 2\"). "
            "Selecting a suggestion populates the field."
        ),
    },
    {
        "id": "TC37",
        "title": "Bottom Navigation — Tab Switching",
        "objective": "Verify bottom navigation tabs switch between pages.",
        "preconditions": "App open.",
        "steps": (
            "1. Tap Home tab.\n"
            "2. Tap Results tab.\n"
            "3. Tap Scoring tab.\n"
            "4. Tap Settings tab."
        ),
        "expected": (
            "Each tab navigates to the corresponding page. Active tab is visually highlighted. "
            "State is preserved when switching back."
        ),
    },
    {
        "id": "TC38",
        "title": "Refresh Route Results",
        "objective": "Verify refreshing re-fetches routes with the same parameters.",
        "preconditions": "A successful search has been performed.",
        "steps": "1. On ResultsPage, tap the Refresh button.",
        "expected": (
            "Spinner appears. Routes are re-fetched from the backend. "
            "Results update (scores/crowding may change if real-time data changed)."
        ),
    },
    {
        "id": "TC39",
        "title": "Google Maps Passthrough Endpoint",
        "objective": "Verify GET /gmaps/directions returns raw directions data.",
        "preconditions": "Backend running; Google Maps API key configured.",
        "steps": "Send: GET /gmaps/directions?origin=1.3,103.8&destination=1.35,103.85&mode=transit",
        "expected": "200 OK. Response contains Google Maps directions data (routes, legs, steps).",
    },
    {
        "id": "TC40",
        "title": "Concurrent Constraint Combination",
        "objective": "Verify multiple constraints applied simultaneously work correctly.",
        "preconditions": "Backend running.",
        "steps": (
            "1. Set Max Walk = 10 min, Max Transfers = 1, Max Budget = $2.00.\n"
            "2. Search for a route."
        ),
        "expected": (
            "Only routes satisfying ALL three constraints are returned. "
            "If no routes match all constraints, \"No routes found\" message is shown."
        ),
    },
]

# -- Group labels --
groups = [
    ("Route Search", 0, 6),
    ("Advanced Priorities & Constraints", 7, 10),
    ("Results & Map", 11, 12),
    ("Scoring", 13, 14),
    ("Departure Time Comparison", 15, 16),
    ("Crowding Heatmap", 17, 17),
    ("Settings", 18, 20),
    ("Risk Badges", 21, 21),
    ("Backend API Endpoints", 22, 28),
    ("Scoring Logic & Fallback", 29, 30),
    ("Feature-Specific", 31, 34),
    ("UI Interactions", 35, 39),
]

HEADER_COLOR = "1F3864"
HEADER_FONT_COLOR = RGBColor(0xFF, 0xFF, 0xFF)
ALT_ROW_COLOR = "F2F2F2"
BORDER_COLOR = "BFBFBF"


def set_cell_shading(cell, color):
    shading = cell._element.get_or_add_tcPr()
    sh = shading.makeelement(qn("w:shd"), {
        qn("w:val"): "clear",
        qn("w:color"): "auto",
        qn("w:fill"): color,
    })
    shading.append(sh)


def add_tc_table(tc):
    rows_data = [
        ("Test Case ID", tc["id"]),
        ("Title", tc["title"]),
        ("Objective", tc["objective"]),
        ("Preconditions", tc["preconditions"]),
        ("Steps", tc["steps"]),
        ("Expected Result", tc["expected"]),
    ]
    table = doc.add_table(rows=len(rows_data), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    for i, (label, value) in enumerate(rows_data):
        row = table.rows[i]
        # Label cell
        cell_l = row.cells[0]
        cell_l.width = Inches(1.4)
        p = cell_l.paragraphs[0]
        run = p.add_run(label)
        run.bold = True
        run.font.size = Pt(9)
        run.font.name = "Calibri"
        set_cell_shading(cell_l, "E8EDF5")

        # Value cell
        cell_r = row.cells[1]
        cell_r.width = Inches(5.4)
        p = cell_r.paragraphs[0]
        run = p.add_run(value)
        run.font.size = Pt(9)
        run.font.name = "Calibri"

    doc.add_paragraph()  # spacer


# -- Build document --
for group_name, start, end in groups:
    doc.add_heading(group_name, level=1)
    for tc in test_cases[start : end + 1]:
        add_tc_table(tc)

# -- Summary table --
doc.add_heading("Summary", level=1)
summary_table = doc.add_table(rows=len(test_cases) + 1, cols=4)
summary_table.style = "Table Grid"
summary_table.alignment = WD_TABLE_ALIGNMENT.CENTER

headers = ["ID", "Title", "Category", "Result"]
for i, h in enumerate(headers):
    cell = summary_table.rows[0].cells[i]
    p = cell.paragraphs[0]
    run = p.add_run(h)
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = HEADER_FONT_COLOR
    set_cell_shading(cell, HEADER_COLOR)

# Map TC index to group
def get_group(idx):
    for gn, s, e in groups:
        if s <= idx <= e:
            return gn
    return ""

for i, tc in enumerate(test_cases):
    row = summary_table.rows[i + 1]
    values = [tc["id"], tc["title"], get_group(i), ""]
    for j, v in enumerate(values):
        p = row.cells[j].paragraphs[0]
        run = p.add_run(v)
        run.font.size = Pt(8)
        run.font.name = "Calibri"
    if i % 2 == 1:
        for j in range(4):
            set_cell_shading(row.cells[j], ALT_ROW_COLOR)

# Set column widths for summary
for row in summary_table.rows:
    row.cells[0].width = Inches(0.5)
    row.cells[1].width = Inches(3.2)
    row.cells[2].width = Inches(2.0)
    row.cells[3].width = Inches(0.8)

out_path = r"C:\Users\teren\Downloads\SC2006-ACDA-G36-main\SC2006-ACDA-G36-main\Program\tests\TestCases.docx"
doc.save(out_path)
print(f"Saved to {out_path}")
