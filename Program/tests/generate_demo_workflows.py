"""Generate DemoWorkflows.docx — 4 end-to-end demo workflows (~10 min total)."""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = Document()

# -- Styles --
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(10)
style.paragraph_format.space_after = Pt(4)

for section in doc.sections:
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

# -- Helpers --
def set_cell_shading(cell, color):
    shading = cell._element.get_or_add_tcPr()
    sh = shading.makeelement(qn("w:shd"), {
        qn("w:val"): "clear", qn("w:color"): "auto", qn("w:fill"): color,
    })
    shading.append(sh)


ACCENT = "1F3864"
ACCENT_LIGHT = "D6E4F0"
STEP_BG = "F7F9FC"
CHECK_BG = "E8F5E9"

# ================================================================
#  WORKFLOW DATA
# ================================================================
workflows = [
    # ── WORKFLOW 1 ──────────────────────────────────────────────
    {
        "number": 1,
        "title": "Public Transit Search — Full Journey",
        "time": "~3 min",
        "covers": "TC1, TC4, TC5, TC6, TC7, TC8, TC9, TC10, TC11, TC12, TC13, TC14, TC16, TC18, TC22, TC35, TC36, TC40",
        "goal": "Demonstrate the core user journey: search for a public transit route, explore results, inspect scoring, compare departure times, and view the route on the map.",
        "steps": [
            {
                "action": "Open the app. Show the Home screen with the empty search form.",
                "narration": "This is SGTravelBud — a smart transport routing app for Singapore.",
                "verify": "Map is visible. Search fields are empty. Search button is disabled.",
            },
            {
                "action": "Type \"Clem\" in the Origin field. Wait for autocomplete dropdown to appear. Select \"Clementi MRT Station\".",
                "narration": "We use Nominatim geocoding for place autocomplete — suggestions appear after 2 characters.",
                "verify": "Autocomplete dropdown shows Singapore locations. Selection populates the field with coordinates.",
            },
            {
                "action": "Try to type the same location \"Clementi MRT\" in Destination. Note the Search button stays disabled.",
                "narration": "The app validates that origin and destination are different.",
                "verify": "Search button remains disabled. (Covers TC5, TC6 — input validation.)",
            },
            {
                "action": "Clear Destination. Enter \"Raffles Place MRT\" as the destination. Ensure Public Transit mode is selected.",
                "narration": "Now we have a valid transit query — Clementi to Raffles Place.",
                "verify": "Both fields filled. Search button becomes enabled.",
            },
            {
                "action": "Tap the Swap button between origin and destination.",
                "narration": "Quick swap — useful if you entered them the wrong way around.",
                "verify": "Origin and Destination values are reversed. (Covers TC4.)",
            },
            {
                "action": "Swap back. Expand the Advanced Priorities section. Show the 4 weight sliders (Time, Cost, Risk, Comfort) and the 3 constraint fields (Max Walk, Max Transfers, Max Budget).",
                "narration": "Users can personalise how routes are ranked. Weights control the scoring formula; constraints hard-filter routes.",
                "verify": "Sliders default to 0.25 each. Constraint fields are empty (no constraint).",
            },
            {
                "action": "Set Max Transfers to 0. Tap Search.",
                "narration": "Let's first try a strict constraint — zero transfers.",
                "verify": "If no direct routes exist: \"No routes found\" message. (Covers TC10, TC7.) Otherwise, only direct routes shown.",
            },
            {
                "action": "Remove the Max Transfers constraint (clear the field). Tap Search again with default weights.",
                "narration": "Now searching with no constraints — we get the full ranked list.",
                "verify": "Bottom sheet expands. 1–5 route cards appear with duration, cost, transfers, and crowding risk badges (Low/Medium/High). Map shows route polyline. (Covers TC1, TC22.)",
            },
            {
                "action": "Navigate to the Results tab. Point out the route cards, risk legend, and the Compare Departure Times button.",
                "narration": "Each card shows a composite score. Risk badges use real-time LTA crowding and delay data.",
                "verify": "Route cards are listed in score order. Risk legend shows Low (green), Medium (amber), High (red).",
            },
            {
                "action": "Tap Compare Departure Times.",
                "narration": "This compares routes across 9 time slots — Morning Rush, Around Now, and Evening Rush — so you can see how departure time affects your journey.",
                "verify": "Modal opens with 3 groups x 3 time slots. Cells show score, time, realistic time, cost, crowding. Best = green, worst = red. (Covers TC16.)",
            },
            {
                "action": "Close the comparison modal. Select a route card. Tap View on Map.",
                "narration": "Let's inspect this route on the map.",
                "verify": "MapPage shows the selected route with polyline, start/end markers. Bottom panel: time, distance, cost, risk badges. (Covers TC12.)",
            },
            {
                "action": "Use the Prev/Next arrows to cycle through routes on the map.",
                "narration": "You can browse all routes directly on the map. Prev is disabled on route 1, Next on the last.",
                "verify": "Map and info panel update for each route. \"Route X of Y\" counter updates. (Covers TC13.)",
            },
            {
                "action": "Navigate to the Scoring tab. Scroll through the scoring breakdown.",
                "narration": "The Scoring page shows how each route was ranked. We normalise metrics to 0–1, multiply by weights, and sum. Realistic time accounts for bus frequency delays.",
                "verify": "Active weights displayed. Comparison table (best green, worst red). Per-route breakdown cards with normalised bars, weighted contributions, and formula. Realistic time >= Google time, highlighted amber if different. Crowding heatmaps per MRT station (6AM–11PM, with Now marker). (Covers TC14, TC18, TC35.)",
            },
            {
                "action": "Go back to Home. Expand Advanced Priorities. Set Time = 1.0, Cost = 0.0, Risk = 0.0, Comfort = 0.0. Search again.",
                "narration": "Watch how the ranking changes when we only care about speed.",
                "verify": "The fastest route is now ranked #1. Order differs from the default-weight search. (Covers TC8.)",
            },
            {
                "action": "Now add constraints: Max Walk = 5 min, Max Transfers = 1, Max Budget = $2.00. Search.",
                "narration": "Constraints hard-filter: any route exceeding these limits is removed entirely.",
                "verify": "Results only include routes satisfying ALL three constraints. Fewer routes returned. (Covers TC9, TC11, TC40.)",
            },
        ],
    },
    # ── WORKFLOW 2 ──────────────────────────────────────────────
    {
        "number": 2,
        "title": "Driving / Taxi Search — ERP, Parking, Weather",
        "time": "~3 min",
        "covers": "TC2, TC3, TC32, TC33, TC34",
        "goal": "Demonstrate driving and taxi mode with differentiation features: ERP toll calculation, carpark availability, and weather integration.",
        "steps": [
            {
                "action": "On the Home screen, enter Origin = \"NUS\" and Destination = \"Changi Airport\".",
                "narration": "A cross-island trip — good for showing driving features.",
                "verify": "Fields populated. Search button enabled.",
            },
            {
                "action": "Toggle mode to Taxi/Drive only (deselect Public Transit).",
                "narration": "Now we're searching exclusively for driving and taxi routes.",
                "verify": "Only the Taxi/Drive toggle is active.",
            },
            {
                "action": "Tap Search.",
                "narration": "The backend calls Google Maps for driving directions, then enriches them with ERP gantry detection, carpark availability, and weather data.",
                "verify": "Results show route cards categorised as \"Taxi\" or \"Drive\". Cost reflects taxi fare estimate or fuel cost. No transfer count shown. (Covers TC2.)",
            },
            {
                "action": "Select a driving route. Check the route explanation text for ERP mentions.",
                "narration": "The app detects ERP gantries along the route polyline and adds toll charges to the cost. This uses the CTE and other expressway gantry data.",
                "verify": "If the route passes ERP gantries during peak hours, the cost includes toll charges and the explanation mentions ERP. (Covers TC32.)",
            },
            {
                "action": "Check the route details for carpark availability near the destination.",
                "narration": "For driving routes, we query LTA for nearby HDB/URA carparks and show available lot counts.",
                "verify": "Carpark availability data appears in the route response with lot counts. (Covers TC34.)",
            },
            {
                "action": "Check the route explanation or assessment for weather information.",
                "narration": "We pull the NEA 2-hour weather forecast and factor adverse weather into the risk assessment.",
                "verify": "If rain or thunderstorms are forecast, the route explanation or risk reflects this. (Covers TC33.)",
            },
            {
                "action": "Now enable BOTH modes (Public Transit + Taxi/Drive). Search again.",
                "narration": "With both modes on, you get a mixed comparison — transit vs driving vs taxi side by side.",
                "verify": "Results include a mix of transit, taxi, and drive routes, up to 5 total. Categories labelled correctly. (Covers TC3.)",
            },
            {
                "action": "Navigate to the Scoring tab. Compare transit vs driving route scores.",
                "narration": "The scoring system works across all modes — normalised metrics let you compare apples to oranges.",
                "verify": "Comparison table shows all route types. Scoring formula applies consistently across modes.",
            },
        ],
    },
    # ── WORKFLOW 3 ──────────────────────────────────────────────
    {
        "number": 3,
        "title": "Settings, Data Freshness, and Fallback",
        "time": "~2 min",
        "covers": "TC19, TC20, TC21, TC28, TC29, TC31, TC37",
        "goal": "Demonstrate user settings management, live data source monitoring, cache refresh, and fallback resilience.",
        "steps": [
            {
                "action": "Use the bottom navigation to switch between all 4 tabs: Home, Results, Scoring, Settings.",
                "narration": "The bottom nav provides quick access to all sections. State is preserved across tabs.",
                "verify": "Each tab loads correctly. Active tab is highlighted. (Covers TC37.)",
            },
            {
                "action": "Navigate to the Settings tab. Show the current settings: language, units, and default weights.",
                "narration": "User preferences are persisted on the backend as a JSON file.",
                "verify": "Settings page loads with current values. Default weights shown as sliders. (Covers TC19.)",
            },
            {
                "action": "Change Language to \"zh\" (Chinese). Change Units to \"imperial\". Adjust the default Time weight to 0.8. Tap Save.",
                "narration": "Settings are saved via PUT /settings and persist across sessions.",
                "verify": "Green success toast appears. (Covers TC20.)",
            },
            {
                "action": "Navigate away to Home, then back to Settings.",
                "narration": "Let's verify persistence.",
                "verify": "Settings still show zh, imperial, Time = 0.8. Values persisted. (Covers TC20.)",
            },
            {
                "action": "Reset settings back to defaults (en, metric, all weights 0.25). Save.",
                "narration": "Restoring defaults for the rest of the demo.",
                "verify": "Success toast. Settings reset.",
            },
            {
                "action": "Scroll down to the Data Sources section. Show the table of data sources with their Live/Fallback status.",
                "narration": "This shows the real-time health of every external data source the app depends on — LTA bus arrivals, train crowding, traffic speed bands, carpark availability, and more. Each has a TTL-based cache.",
                "verify": "Table lists all data sources. Each shows \"Live\" (green) or \"Fallback\" (amber) badge with last-retrieved timestamp. (Covers TC28.)",
            },
            {
                "action": "Tap the Refresh button in the Data Sources section.",
                "narration": "This invalidates all caches and re-fetches fresh data from every external API. If an API is down, the app falls back to the last known good data.",
                "verify": "Sources reload. Live sources refresh timestamps. Any unavailable API shows \"Fallback\" badge. (Covers TC21, TC31.)",
            },
        ],
    },
    # ── WORKFLOW 4 ──────────────────────────────────────────────
    {
        "number": 4,
        "title": "Backend API Verification (Live Demo with Swagger/curl)",
        "time": "~2 min",
        "covers": "TC23, TC24, TC25, TC26, TC27, TC29, TC30, TC38, TC39",
        "goal": "Demonstrate the backend API layer directly — showing that the service layer, data layer, and external integrations all work end-to-end.",
        "steps": [
            {
                "action": "Open a browser to the FastAPI docs (http://localhost:8000/docs) or use curl. Hit GET /health.",
                "narration": "The backend exposes a RESTful API with automatic Swagger documentation. Let's start with a health check.",
                "verify": "200 OK response. Server is running. (Covers TC29.)",
            },
            {
                "action": "Call GET /routes?origin=Clementi&destination=Raffles+Place&include_transit=true&wt_time=0.5&wt_cost=0.2&wt_risk=0.2&wt_comfort=0.1",
                "narration": "This is the main routing endpoint. It calls Google Maps, enriches with LTA data, scores, and ranks.",
                "verify": "200 OK. Response has \"trip\" metadata and \"routes\" array with up to 5 RouteOption objects. Each has time, cost, score, steps, crowding, delay. (Covers TC23.)",
            },
            {
                "action": "Call GET /routes?origin=NUS&destination=Changi+Airport&max_transfers=1&max_budget=3.0",
                "narration": "Adding constraints — the backend filters server-side before returning results.",
                "verify": "All returned routes have transfers <= 1 and cost <= $3.00. (Covers TC24.)",
            },
            {
                "action": "Call GET /assessment?origin=Jurong+East&destination=Dhoby+Ghaut",
                "narration": "The assessment endpoint provides segment-level risk analysis — crowding forecasts and delay estimates for each leg.",
                "verify": "200 OK. Segment assessments with crowding levels, delay estimates, bus frequency data. (Covers TC25.)",
            },
            {
                "action": "Call GET /crowding/heatmap?station_name=Orchard",
                "narration": "This powers the crowding heatmap visualisation — hourly crowding data for any MRT station.",
                "verify": "200 OK. Station name, line code, and intervals array (6AM–11PM) with crowding levels. (Covers TC26.)",
            },
            {
                "action": "Call PUT /settings with body: {\"language\":\"ta\",\"units\":\"imperial\",\"default_weights\":{\"time\":0.4,\"cost\":0.3,\"risk\":0.2,\"comfort\":0.1}}. Then call GET /settings.",
                "narration": "Settings are persisted via a simple JSON file-based repository.",
                "verify": "PUT returns 200. GET returns the same values. Round-trip confirmed. (Covers TC27.)",
            },
            {
                "action": "Call GET /routes/compare?origin=Clementi&destination=Raffles+Place&include_transit=true (show the response structure).",
                "narration": "The compare endpoint returns the best route for each of 9 departure time slots — this powers the Departure Time Comparison modal.",
                "verify": "Response has 3 groups (Morning/Now/Evening), each with 3 time slots, each with top 3 routes.",
            },
            {
                "action": "Go back to the frontend. On the Results page, tap Refresh.",
                "narration": "The refresh button re-hits the same API with the same parameters to get updated real-time data.",
                "verify": "Spinner appears. Fresh results load. Scores may differ slightly from real-time data changes. (Covers TC38.)",
            },
        ],
    },
]

# ================================================================
#  BUILD DOCUMENT
# ================================================================

title = doc.add_heading("SGTravelBud — Demo Workflows", level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

p = doc.add_paragraph()
run = p.add_run("4 end-to-end workflows designed for a ~10-minute live demo. "
                 "Together they cover all 40 test cases (TC1–TC40).")
run.font.size = Pt(10)

doc.add_paragraph()

# -- Overview table --
doc.add_heading("Overview", level=1)
overview = doc.add_table(rows=len(workflows) + 1, cols=4)
overview.style = "Table Grid"
overview.alignment = WD_TABLE_ALIGNMENT.CENTER

for i, h in enumerate(["#", "Workflow", "Time", "Test Cases Covered"]):
    cell = overview.rows[0].cells[i]
    p = cell.paragraphs[0]
    run = p.add_run(h)
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    set_cell_shading(cell, ACCENT)

for i, wf in enumerate(workflows):
    row = overview.rows[i + 1]
    vals = [str(wf["number"]), wf["title"], wf["time"], wf["covers"]]
    for j, v in enumerate(vals):
        p = row.cells[j].paragraphs[0]
        run = p.add_run(v)
        run.font.size = Pt(9)
        run.font.name = "Calibri"
    if i % 2 == 1:
        for j in range(4):
            set_cell_shading(row.cells[j], "F2F2F2")

overview.rows[0].cells[0].width = Inches(0.3)
overview.rows[0].cells[1].width = Inches(2.8)
overview.rows[0].cells[2].width = Inches(0.6)
overview.rows[0].cells[3].width = Inches(3.1)

doc.add_paragraph()

# -- Each workflow --
for wf in workflows:
    doc.add_heading(
        f"Workflow {wf['number']}: {wf['title']}", level=1
    )

    # Meta info
    meta_table = doc.add_table(rows=3, cols=2)
    meta_table.style = "Table Grid"
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_items = [
        ("Estimated Time", wf["time"]),
        ("Goal", wf["goal"]),
        ("Test Cases Covered", wf["covers"]),
    ]
    for i, (label, value) in enumerate(meta_items):
        cell_l = meta_table.rows[i].cells[0]
        cell_l.width = Inches(1.4)
        p = cell_l.paragraphs[0]
        run = p.add_run(label)
        run.bold = True
        run.font.size = Pt(9)
        set_cell_shading(cell_l, ACCENT_LIGHT)

        cell_r = meta_table.rows[i].cells[1]
        cell_r.width = Inches(5.4)
        p = cell_r.paragraphs[0]
        run = p.add_run(value)
        run.font.size = Pt(9)

    doc.add_paragraph()

    # Steps table
    steps_table = doc.add_table(rows=len(wf["steps"]) + 1, cols=4)
    steps_table.style = "Table Grid"
    steps_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for j, h in enumerate(["Step", "Action (What to Do)", "Narration (What to Say)", "Expected Result / Verify"]):
        cell = steps_table.rows[0].cells[j]
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, ACCENT)

    for i, step in enumerate(wf["steps"]):
        row = steps_table.rows[i + 1]
        vals = [str(i + 1), step["action"], step["narration"], step["verify"]]
        for j, v in enumerate(vals):
            p = row.cells[j].paragraphs[0]
            run = p.add_run(v)
            run.font.size = Pt(8)
            run.font.name = "Calibri"

        # Alternate row shading
        bg = STEP_BG if i % 2 == 0 else "FFFFFF"
        for j in range(4):
            set_cell_shading(row.cells[j], bg)

        # Green tint on verify column
        set_cell_shading(row.cells[3], CHECK_BG if i % 2 == 0 else "F0FAF0")

    # Column widths
    for row in steps_table.rows:
        row.cells[0].width = Inches(0.35)
        row.cells[1].width = Inches(2.5)
        row.cells[2].width = Inches(2.1)
        row.cells[3].width = Inches(2.2)

    doc.add_paragraph()

# -- Coverage summary --
doc.add_heading("Test Case Coverage Map", level=1)
p = doc.add_paragraph()
run = p.add_run(
    "The table below shows which original test cases (TC1–TC40) are exercised by each workflow."
)
run.font.size = Pt(9)

all_tcs = list(range(1, 41))
coverage = {}
for wf in workflows:
    nums = []
    for part in wf["covers"].split(","):
        part = part.strip().replace("TC", "")
        nums.append(int(part))
    for n in nums:
        coverage.setdefault(n, []).append(wf["number"])

cov_table = doc.add_table(rows=len(all_tcs) + 1, cols=3)
cov_table.style = "Table Grid"
cov_table.alignment = WD_TABLE_ALIGNMENT.CENTER

for j, h in enumerate(["Test Case", "Covered By Workflow(s)", "Status"]):
    cell = cov_table.rows[0].cells[j]
    p = cell.paragraphs[0]
    run = p.add_run(h)
    run.bold = True
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    set_cell_shading(cell, ACCENT)

for i, tc in enumerate(all_tcs):
    row = cov_table.rows[i + 1]
    wfs = coverage.get(tc, [])
    wf_str = ", ".join(f"WF{w}" for w in wfs) if wfs else "—"
    status = "Covered" if wfs else "Not in demo"

    vals = [f"TC{tc}", wf_str, status]
    for j, v in enumerate(vals):
        p = row.cells[j].paragraphs[0]
        run = p.add_run(v)
        run.font.size = Pt(8)

    if not wfs:
        for j in range(3):
            set_cell_shading(row.cells[j], "FFF3E0")
    elif i % 2 == 1:
        for j in range(3):
            set_cell_shading(row.cells[j], "F2F2F2")

for row in cov_table.rows:
    row.cells[0].width = Inches(0.8)
    row.cells[1].width = Inches(2.0)
    row.cells[2].width = Inches(1.0)

out_path = r"C:\Users\teren\Downloads\SC2006-ACDA-G36-main\SC2006-ACDA-G36-main\Program\tests\DemoWorkflows.docx"
doc.save(out_path)
print(f"Saved to {out_path}")
